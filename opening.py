from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def open(d):
    f = Document(d)
    t = []
    for i in f.paragraphs:
        t.append(i.text)
    return t

doc = Document()
rights = []

def add(str):
    doc.add_paragraph(str)

def addcenter(str):
    pr = doc.add_paragraph(str)
    pr.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

def writerights():
    print(rights)
    varnow = rights[0][0]
    answers = "Вар 1. "
    for list in rights:
        if list[0] != varnow:
            varnow = list[0]
            answers += f"Вар {varnow}. "
        answers += f"{list[1]} - {list[2]}), "
    doc.add_paragraph(answers[:len(answers)-2])

def addright(qes, ans, var):
    rights.append([var, qes, ans])

def new_page():
    doc.add_page_break()

def savedoc(head):
    writerights()
    name = head+".docx"
    doc.save(name)
    os.startfile(name)
